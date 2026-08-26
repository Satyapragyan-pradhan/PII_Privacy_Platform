import io

from docx import Document


def extract_docx_text(content: bytes):
    document = Document(
        io.BytesIO(content)
    )

    paragraphs = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()

        if text:
            paragraphs.append(text)

    # Extract table contents too
    for table in document.tables:
        for row in table.rows:
            row_text = []

            for cell in row.cells:
                text = cell.text.strip()

                if text:
                    row_text.append(text)

            if row_text:
                paragraphs.append(" | ".join(row_text))

    return {
        "text": "\n".join(paragraphs),
        "needs_ocr": False,
        "pages": []
    }